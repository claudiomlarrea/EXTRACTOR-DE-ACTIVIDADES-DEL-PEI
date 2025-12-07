import streamlit as st
import pandas as pd
import numpy as np
import re
from io import BytesIO

from unidecode import unidecode
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document


# -----------------------------
# Utilidades de normalización
# -----------------------------

def normalize_text(s: str) -> str:
    """Normaliza texto: minúsculas, sin tildes, espacios simples."""
    if pd.isna(s):
        return ""
    s = str(s)
    s = unidecode(s)          # quita tildes
    s = s.lower().strip()
    s = re.sub(r"\s+", " ", s)
    return s


def normalize_colnames(df: pd.DataFrame) -> dict:
    """Devuelve dict nombre_normalizado -> nombre_original."""
    mapping = {}
    for col in df.columns:
        norm = normalize_text(col)
        mapping[norm] = col
    return mapping


def find_column(norm_map: dict, patterns) -> str | None:
    """
    Busca la primera columna cuyo nombre normalizado contenga
    alguno de los patrones indicados.
    """
    if isinstance(patterns, str):
        patterns = [patterns]

    for norm_name, original in norm_map.items():
        for pat in patterns:
            if pat in norm_name:
                return original
    return None


# -----------------------------
# Extracción de actividades únicas desde el Formulario Único
# -----------------------------

def extract_activities_from_form(df: pd.DataFrame) -> pd.DataFrame:
    """
    A partir del Formulario Único (como sale de Looker / Google Sheets),
    genera una tabla “plana” de actividades únicas:

    ID, Año, Objetivo, Actividad, Detalle
    """

    norm_map = normalize_colnames(df)

    # Columna Año
    year_col = find_column(
        norm_map,
        ["ano", "año"]  # sin tilde gracias a unidecode
    )

    # Si no hay ID, generamos uno incremental
    id_col = find_column(norm_map, ["id "])  # ej. "id " o "id actividad"
    if id_col is None:
        df = df.copy()
        df["ID"] = np.arange(1, len(df) + 1)
        id_col = "ID"

    records = []

    # Buscamos hasta 6 objetivos / actividades / detalles (1..6)
    for i in range(1, 7):
        # Objetivo específico i
        obj_col = find_column(
            norm_map,
            [
                f"objetivos especificos {i}",
                f"objetivo especifico {i}",
                f"objetivo {i}"
            ]
        )
        # Actividad objetivo i
        act_col = find_column(
            norm_map,
            [
                f"actividades objetivo {i}",
                f"actividad objetivo {i}",
                f"actividad obj {i}"
            ]
        )
        # Detalle actividad objetivo i
        det_col = find_column(
            norm_map,
            [
                f"detalle de la actividad objetivo {i}",
                f"detalle actividad objetivo {i}",
                f"detalle obj {i}"
            ]
        )

        # Si no hay columna de objetivo o actividad, pasamos al siguiente i
        if obj_col is None or act_col is None:
            continue

        for _, row in df.iterrows():
            obj = row.get(obj_col, "")
            act = row.get(act_col, "")
            det = row.get(det_col, "") if det_col else ""

            if (pd.isna(obj) or str(obj).strip() == "") and \
               (pd.isna(act) or str(act).strip() == ""):
                continue  # no hay nada para este objetivo en esta fila

            record = {
                "ID": row[id_col],
                "Año": row.get(year_col, None),
                "Objetivo específico": obj,
                "Actividad única": act,
                "Detalle actividad": det
            }
            records.append(record)

    activities_df = pd.DataFrame(records)

    # Limpieza básica
    if not activities_df.empty:
        activities_df["Año"] = activities_df["Año"].astype(str).str.strip()
        activities_df["Objetivo específico"] = activities_df["Objetivo específico"].astype(str).str.strip()
        activities_df["Actividad única"] = activities_df["Actividad única"].astype(str).str.strip()
        activities_df["Detalle actividad"] = activities_df["Detalle actividad"].astype(str).str.strip()

    return activities_df


# -----------------------------
# Cálculo de consistencia
# -----------------------------

def map_similarity_to_score(sim: float) -> int:
    """
    Mapea la similitud coseno (0-1) a los niveles discretos:
    0, 10, 30, 50, 70, 90, 100.
    """
    if np.isnan(sim) or sim <= 0:
        return 0
    if sim < 0.10:
        return 0
    elif sim < 0.25:
        return 10
    elif sim < 0.40:
        return 30
    elif sim < 0.55:
        return 50
    elif sim < 0.70:
        return 70
    elif sim < 0.90:
        return 90
    else:
        return 100


def compute_consistency(activities_df: pd.DataFrame) -> pd.DataFrame:
    """
    Calcula la similitud semántica entre Objetivo específico y Actividad única
    usando TF-IDF + coseno, y añade columnas:
        - Similitud
        - Consistencia (%)
    """

    df = activities_df.copy()

    objetivos = df["Objetivo específico"].fillna("").astype(str)
    actividades = (
        df["Actividad única"].fillna("").astype(str) + " " +
        df["Detalle actividad"].fillna("").astype(str)
    )

    # Corpus conjunto para vectorizar
    corpus = objetivos.tolist() + actividades.tolist()

    vectorizer = TfidfVectorizer(
        ngram_range=(1, 2),
        stop_words="spanish"
    )
    vectorizer.fit(corpus)

    obj_vecs = vectorizer.transform(objetivos)
    act_vecs = vectorizer.transform(actividades)

    sims = []
    for i in range(obj_vecs.shape[0]):
        sim = cosine_similarity(obj_vecs[i], act_vecs[i])[0, 0]
        sims.append(sim)

    df["Similitud"] = sims
    df["Consistencia (%)"] = df["Similitud"].apply(map_similarity_to_score)

    return df


# -----------------------------
# Generación de Excel
# -----------------------------

def build_summary_sheets(df_cons: pd.DataFrame):
    """
    Construye:
      - resumen por Año y Objetivo
      - indicadores globales
      - distribución de actividades por nivel de consistencia
    """
    # Resumen Año / Objetivo
    resumen = (
        df_cons
        .groupby(["Año", "Objetivo específico"], dropna=False)
        .agg(
            Cant_actividades=("Actividad única", "count"),
            Consistencia_promedio=("Consistencia (%)", "mean")
        )
        .reset_index()
    )
    resumen["Consistencia_promedio"] = resumen["Consistencia_promedio"].round(2)

    # Indicadores globales
    total_acts = len(df_cons)
    consist_prom = df_cons["Consistencia (%)"].mean() if total_acts > 0 else 0
    consist_prom = round(consist_prom, 2)

    indicadores = pd.DataFrame({
        "Indicador": [
            "Cantidad total de actividades únicas",
            "Consistencia general (%)"
        ],
        "Valor": [
            total_acts,
            consist_prom
        ]
    })

    # Distribución de niveles de consistencia
    dist = (
        df_cons["Consistencia (%)"]
        .value_counts()
        .sort_index()
        .rename_axis("Consistencia (%)")
        .reset_index(name="Cantidad")
    )

    return resumen, indicadores, dist


def generate_excel_bytes(df_cons: pd.DataFrame,
                         resumen: pd.DataFrame,
                         indicadores: pd.DataFrame,
                         dist: pd.DataFrame) -> bytes:
    """Genera un Excel con varias hojas y lo devuelve como bytes."""

    output = BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df_cons.to_excel(writer, sheet_name="Detalle actividades", index=False)
        resumen.to_excel(writer, sheet_name="Resumen año-objetivo", index=False)
        indicadores.to_excel(writer, sheet_name="Indicadores globales", index=False)
        dist.to_excel(writer, sheet_name="Distribución consistencia", index=False)

    output.seek(0)
    return output.read()


# -----------------------------
# Generación de informe Word
# -----------------------------

def generate_word_report_bytes(indicadores: pd.DataFrame,
                               dist: pd.DataFrame) -> bytes:
    """
    Construye un informe en Word con:
      - indicadores globales
      - análisis breve
      - tabla de distribución por niveles de consistencia
    """

    doc = Document()
    doc.add_heading("Informe de Consistencia PEI", level=1)

    # Indicadores globales
    doc.add_heading("1. Indicadores globales", level=2)

    total_acts = int(indicadores.loc[indicadores["Indicador"] ==
                                     "Cantidad total de actividades únicas", "Valor"].values[0])
    consist_gen = float(indicadores.loc[indicadores["Indicador"] ==
                                        "Consistencia general (%)", "Valor"].values[0])

    p = doc.add_paragraph()
    p.add_run("Cantidad total de actividades únicas: ").bold = True
    p.add_run(str(total_acts))

    p = doc.add_paragraph()
    p.add_run("Consistencia general promedio: ").bold = True
    p.add_run(f"{consist_gen:.2f} %")

    # Distribución
    doc.add_heading("2. Distribución por niveles de consistencia", level=2)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Consistencia (%)"
    hdr_cells[1].text = "Cantidad de actividades"

    for _, row in dist.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(int(row["Consistencia (%)"]))
        row_cells[1].text = str(int(row["Cantidad"]))

    # Análisis e interpretación
    doc.add_heading("3. Análisis e interpretación", level=2)

    analysis_par = doc.add_paragraph()
    analysis_par.add_run(
        "El indicador de consistencia general refleja el grado de concordancia "
        "entre los objetivos específicos declarados por las unidades académicas "
        "y las actividades únicas que informan como acciones de cumplimiento del PEI. "
    )

    if consist_gen < 30:
        doc.add_paragraph(
            "El valor promedio obtenido se encuentra en un rango bajo. "
            "Esto sugiere que muchas actividades podrían estar siendo cargadas "
            "en objetivos que no se corresponden plenamente con su finalidad. "
            "Se recomienda revisar la redacción de los objetivos, ofrecer "
            "orientaciones más precisas para la carga de actividades y realizar "
            "instancias de capacitación focalizada."
        )
    elif consist_gen < 60:
        doc.add_paragraph(
            "El valor promedio de consistencia se ubica en un rango medio. "
            "Existe una relación razonable entre objetivos y actividades, "
            "aunque persisten desajustes que pueden corregirse mediante "
            "instancias de retroalimentación con las unidades y ajustes "
            "en la planificación operativa."
        )
    else:
        doc.add_paragraph(
            "El valor promedio de consistencia se encuentra en un rango alto. "
            "En términos generales, las actividades están bien alineadas con "
            "los objetivos específicos planteados en el PEI. "
            "Es conveniente consolidar estas buenas prácticas y focalizar "
            "las acciones de mejora en aquellos objetivos con menor consistencia."
        )

    doc.add_heading("4. Conclusiones", level=2)
    doc.add_paragraph(
        "El análisis de consistencia entre objetivos específicos y actividades únicas "
        "permite disponer de un indicador sintético de calidad de la planificación. "
        "Este insumo puede incorporarse a los procesos de monitoreo institucional, "
        "retroalimentando la formulación de proyectos, la asignación de recursos "
        "y la toma de decisiones estratégicas en cada unidad académica."
    )

    # Exportar a bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# -----------------------------
# Interfaz Streamlit
# -----------------------------

st.set_page_config(
    page_title="Consistencia PEI: Objetivos vs Actividades",
    layout="wide"
)

st.title("Análisis de consistencia entre Objetivos Específicos y Actividades Únicas")
st.write(
    "Subí el archivo Excel del **Formulario Único del PEI**. "
    "La aplicación extraerá todas las actividades únicas, calculará la "
    "consistencia semántica entre cada objetivo específico y la actividad "
    "cargada, y generará un Excel y un informe en Word."
)

uploaded_file = st.file_uploader(
    "Subir archivo Excel del Formulario Único",
    type=["xlsx", "xls"]
)

if uploaded_file is not None:
    try:
        raw_df = pd.read_excel(uploaded_file)
    except Exception as e:
        st.error(f"Error al leer el archivo: {e}")
        st.stop()

    st.subheader("Vista previa del archivo original")
    st.dataframe(raw_df.head())

    # 1) Extraer actividades únicas
    activities_df = extract_activities_from_form(raw_df)

    if activities_df.empty:
        st.error(
            "No se pudieron extraer actividades únicas. "
            "Revisá que el archivo tenga columnas de 'Objetivos específicos', "
            "'Actividades Objetivo' y 'Detalle de la Actividad Objetivo'."
        )
        st.stop()

    st.subheader("Actividades únicas extraídas")
    st.write(f"Se extrajeron **{len(activities_df)}** actividades.")
    st.dataframe(activities_df.head())

    # 2) Calcular consistencia
    df_cons = compute_consistency(activities_df)

    resumen, indicadores, dist = build_summary_sheets(df_cons)

    st.subheader("Indicadores globales")
    st.table(indicadores)

    st.subheader("Distribución de actividades por nivel de consistencia")
    st.table(dist)

    st.subheader("Detalle de actividades con consistencia")
    st.dataframe(df_cons)

    # 3) Descarga de Excel
    excel_bytes = generate_excel_bytes(df_cons, resumen, indicadores, dist)
    st.download_button(
        label="📥 Descargar resultados en Excel",
        data=excel_bytes,
        file_name="consistencia_pei_objetivo_actividad.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

    # 4) Descarga de informe Word
    word_bytes = generate_word_report_bytes(indicadores, dist)
    st.download_button(
        label="📄 Descargar informe en Word",
        data=word_bytes,
        file_name="informe_consistencia_pei.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
